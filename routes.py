import os
import tempfile
from flask import Blueprint, render_template, request, jsonify, redirect, url_for, flash, send_file, session, current_app,abort
from flask_login import login_user, logout_user, login_required, current_user
from models import User, Lesson, Presentation, RoleConfig, TokenTransaction
from forms import LoginForm, RegistrationForm, LessonForm, EditLessonForm, ARLessonForm, UserProfileForm, WhatsAppMessageForm
from lesson_generator import generate_lesson_plan, gpt_plans
from docx import Document
# from ppt_generator import create_presentation
# from whatsapp_sender import process_excel_file, open_whatsapp_web
from urllib.parse import urlparse

routes = Blueprint("routes", __name__)

@routes.route('/')
def index():
    if current_user.is_authenticated:
        return redirect(url_for('routes.dashboard'))
    language = session.get('language', 'en')
    return render_template('index_new.html', language=language)

@routes.route('/api/register', methods=['POST'])
def register_api():
    data = request.get_json()

    # Check if user already exists using Supabase
    if User.get_by_email(data.get('email')):
        return jsonify({'success': False, 'message': 'Email already registered'}), 400

    # Create new user using Supabase
    user = User.create(
        name=data.get('name'),
        email=data.get('email'),
        password=data.get('password')
    )

    if not user:
        return jsonify({'success': False, 'message': 'Registration failed'}), 500

    return jsonify({'success': True, 'message': 'Registration successful'})

@routes.route('/api/login', methods=['POST'])
def login_api():
    data = request.get_json()
    user = User.get_by_email(data.get('email'))

    if user and user.check_password(data.get('password')):
        login_user(user)
        return jsonify({'success': True, 'user': {'id': user.id, 'name': user.name, 'email': user.email}})

    return jsonify({'success': False, 'message': 'Invalid email or password'}), 401

@routes.route('/api/logout', methods=['POST'])
@login_required
def logout_api():
    logout_user()
    return jsonify({'success': True})

@routes.route('/api/user')
@login_required
def get_user():
    return jsonify({
        'id': current_user.id,
        'name': current_user.name,
        'email': current_user.email,
        'role': getattr(current_user, 'role', None),
        'monthly_token_quota': getattr(current_user, 'monthly_token_quota', None),
        'token_balance': getattr(current_user, 'token_balance', None),
        'token_renewal_date': getattr(current_user, 'token_renewal_date', None)
    })

@routes.route('/api/lessons', methods=['GET'])
@login_required
def get_lessons():
    lessons = Lesson.get_all_by_user(current_user.id)
    return jsonify([lesson.to_dict() for lesson in lessons])

@routes.route('/api/lessons', methods=['POST'])
@login_required
def create_lesson():
    data = request.get_json()

    if not current_user.deduct_tokens(1, 'lesson_create', 'app'):
        return jsonify({'success': False, 'message': 'Insufficient tokens', 'notify': 'نفدت التوكنز المتاحة. الرجاء الترقية أو انتظار التجديد الشهري.'}), 403

    # Generate lesson plan content first
    try:
        lesson_plan = generate_lesson_plan(
            data.get('grade_level'),
            data.get('topic'),
            data.get('teaching_strategy'),
            data.get('language')
        )
        gpt_plan = gpt_plans(
            data.get('grade_level'),
            data.get('topic'),
            data.get('teaching_strategy'),
            data.get('language')
        )
    except Exception as e:
        return jsonify({'success': False, 'message': f'Error generating lesson plan: {str(e)}'}), 500

    # Create new lesson using Supabase
    lesson = Lesson.create(
        user_id=current_user.id,
        grade_level=data.get('grade_level'),
        topic=data.get('topic'),
        teaching_strategy=data.get('teaching_strategy'),
        language=data.get('language'),
        generated_plan=lesson_plan,
        gpt_plan=gpt_plan
    )

    if not lesson:
        return jsonify({'success': False, 'message': 'Failed to create lesson'}), 500

    return jsonify({'success': True, 'lesson': lesson.to_dict()})

@routes.route('/api/attendance/confirm', methods=['POST'])
@login_required
def confirm_attendance():
    data = request.get_json() or {}
    amount = int(data.get('amount') or 1)
    reason = data.get('reason') or 'attendance'
    if not current_user.deduct_tokens(amount, reason, 'app'):
        return jsonify({'success': False, 'message': 'Insufficient tokens', 'notify': 'نفدت التوكنز المتاحة. الرجاء الترقية أو انتظار التجديد الشهري.'}), 403
    return jsonify({'success': True, 'balance': current_user.token_balance})

@routes.route('/api/admin/users/<user_id>/reset-balance', methods=['POST'])
@login_required
def admin_reset_balance(user_id):
    require_admin()
    
    target = User.get_by_id(user_id)
    if not target:
        return jsonify({'success': False, 'message': 'User not found'}), 404
    
    # Get the quota for the user's current role
    quota = RoleConfig.get_quota_for_role(target.role)
    
    # Reset balance to quota
    supabase = current_app.config["SUPABASE_CLIENT"]
    supabase.table('users').update({
        'token_balance': quota
    }).eq('id', user_id).execute()
    
    # Optional: Log this action in TokenTransaction
    try:
        TokenTransaction.create(
            user_id=user_id,
            amount=quota,
            transaction_type='admin_reset',
            source='admin',
            reason='Balance reset by admin'
        )
    except:
        pass  # If transaction logging fails, continue anyway
    
    return jsonify({'success': True, 'new_balance': quota})


def require_admin():
    if not current_user.is_admin():
        abort(404)

@routes.route('/admin')
@login_required
def admin_dashboard():
    require_admin()
    language = session.get('language', 'en')
    supabase = current_app.config["SUPABASE_CLIENT"]
    r = supabase.table('users').select("id,name,email,role,monthly_token_quota,token_balance").order('name').execute()
    users = r.data or []
    return render_template('ادمن/dashboard.html', users=users, language=language)

@routes.route('/api/admin/users', methods=['GET'])
@login_required
def admin_list_users():
    require_admin()
    supabase = current_app.config["SUPABASE_CLIENT"]
    r = supabase.table('users').select("id,name,email,role,monthly_token_quota,token_balance").order('name').execute()
    return jsonify({'success': True, 'users': r.data or []})

@routes.route('/api/admin/users/<user_id>/role', methods=['PUT'])
@login_required
def admin_update_role(user_id):
    require_admin()
    data = request.get_json() or {}
    new_role = data.get('role')
    reset_balance = bool(data.get('reset_balance', False))
    target = User.get_by_id(user_id)
    if not target:
        return jsonify({'success': False, 'message': 'User not found'}), 404
    quota = RoleConfig.get_quota_for_role(new_role)
    update = {
        'role': new_role,
        'monthly_token_quota': quota
    }
    if reset_balance:
        update['token_balance'] = quota
    supabase = current_app.config["SUPABASE_CLIENT"]
    supabase.table('users').update(update).eq('id', user_id).execute()
    return jsonify({'success': True})

@routes.route('/api/admin/users/<user_id>/tokens', methods=['PUT'])
@login_required
def admin_update_tokens(user_id):
    require_admin()
    data = request.get_json() or {}
    adjust = data.get('adjust')
    set_quota = data.get('monthly_token_quota')
    set_balance = data.get('token_balance')
    supabase = current_app.config["SUPABASE_CLIENT"]
    target = User.get_by_id(user_id)
    if not target:
        return jsonify({'success': False, 'message': 'User not found'}), 404
    update = {}
    if isinstance(set_quota, int):
        update['monthly_token_quota'] = set_quota
    if isinstance(set_balance, int):
        update['token_balance'] = set_balance
    if update:
        supabase.table('users').update(update).eq('id', user_id).execute()
    if isinstance(adjust, int) and adjust != 0:
        if adjust > 0:
            target.add_tokens(adjust, 'admin_adjustment', 'admin')
        else:
            if not target.deduct_tokens(-adjust, 'admin_adjustment', 'admin'):
                return jsonify({'success': False, 'message': 'Insufficient tokens for deduction'}), 400
    return jsonify({'success': True})

@routes.route('/api/admin/role-configs', methods=['PUT'])
@login_required
def admin_update_role_configs():
    require_admin()
    data = request.get_json() or {}
    role = data.get('role')
    monthly_quota = data.get('monthly_quota')
    if not role or not isinstance(monthly_quota, int):
        return jsonify({'success': False, 'message': 'Invalid data'}), 400
    supabase = current_app.config["SUPABASE_CLIENT"]
    try:
        existing = supabase.table('role_configs').select("*").eq('role', role).execute()
        if existing.data:
            supabase.table('role_configs').update({'monthly_quota': monthly_quota}).eq('role', role).execute()
        else:
            supabase.table('role_configs').insert({'role': role, 'monthly_quota': monthly_quota}).execute()
        return jsonify({'success': True})
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500

@routes.route('/api/admin/role-configs', methods=['GET'])
@login_required
def admin_list_role_configs():
    require_admin()
    supabase = current_app.config["SUPABASE_CLIENT"]
    r = supabase.table('role_configs').select("*").order('role').execute()
    return jsonify({'success': True, 'roles': r.data or []})

@routes.route('/api/lessons/<int:lesson_id>', methods=['GET'])
@login_required
def get_lesson(lesson_id):
    lesson = Lesson.get_by_id(lesson_id)
    
    if not lesson:
        return jsonify({'success': False, 'message': 'Lesson not found'}), 404

    # Check if the lesson belongs to the current user
    if lesson.user_id != current_user.id:
        return jsonify({'success': False, 'message': 'Unauthorized access'}), 403

    return jsonify(lesson.to_dict())

@routes.route('/api/lessons/<int:lesson_id>', methods=['PUT'])
@login_required
def update_lesson(lesson_id):
    lesson = Lesson.get_by_id(lesson_id)
    
    if not lesson:
        return jsonify({'success': False, 'message': 'Lesson not found'}), 404

    # Check if the lesson belongs to the current user
    if lesson.user_id != current_user.id:
        return jsonify({'success': False, 'message': 'Unauthorized access'}), 403

    data = request.get_json()

    # Update lesson using Supabase
    try:
        supabase = current_app.config["SUPABASE_CLIENT"]
        update_data = {}
        
        if 'generated_plan' in data:
            update_data['generated_plan'] = data['generated_plan']
        if 'gpt_plan' in data:
            update_data['gpt_plan'] = data['gpt_plan']
        
        update_data['date_modified'] = 'now()'
        
        response = supabase.table('lessons').update(update_data).eq('id', lesson_id).execute()
        
        # Get updated lesson
        updated_lesson = Lesson.get_by_id(lesson_id)
        return jsonify({'success': True, 'lesson': updated_lesson.to_dict()})
    except Exception as e:
        return jsonify({'success': False, 'message': f'Error updating lesson: {str(e)}'}), 500

@routes.route('/lessons/<int:lesson_id>/delete', methods=['POST'])
@login_required
def delete_lesson_api(lesson_id):
    lesson = Lesson.get_by_id(lesson_id)
    
    if not lesson:
        flash('Lesson not found')
        return redirect(url_for('routes.index'))

    # Check if the lesson belongs to the current user
    if lesson.user_id != current_user.id:
        flash('Unauthorized access')
        return redirect(url_for('routes.index'))

    try:
        supabase = current_app.config["SUPABASE_CLIENT"]
        
        # Delete associated presentations first
        supabase.table('presentations').delete().eq('lesson_id', lesson_id).execute()
        
        # Delete the lesson
        supabase.table('lessons').delete().eq('id', lesson_id).execute()
        
        flash('Lesson plan deleted successfully')
    except Exception as e:
        flash(f'Error deleting lesson: {str(e)}')
    
    return redirect(url_for('routes.get_lessons_page'))

@routes.route('/api/lessons/<int:lesson_id>', methods=['DELETE'])
@login_required
def delete_lesson(lesson_id):
    lesson = Lesson.get_by_id(lesson_id)
    
    if not lesson:
        return jsonify({'success': False, 'message': 'Lesson not found'}), 404

    # Check if the lesson belongs to the current user
    if lesson.user_id != current_user.id:
        return jsonify({'success': False, 'message': 'Unauthorized access'}), 403

    try:
        supabase = current_app.config["SUPABASE_CLIENT"]
        
        # Delete associated presentations first
        supabase.table('presentations').delete().eq('lesson_id', lesson_id).execute()
        
        # Delete the lesson
        supabase.table('lessons').delete().eq('id', lesson_id).execute()
        
        return jsonify({'success': True})
    except Exception as e:
        return jsonify({'success': False, 'message': f'Error deleting lesson: {str(e)}'}), 500

@routes.route('/api/lessons/<int:lesson_id>/presentation', methods=['POST'])
@login_required
def generate_presentation(lesson_id):
    lesson = Lesson.get_by_id(lesson_id)
    
    if not lesson:
        return jsonify({'success': False, 'message': 'Lesson not found'}), 404

    # Check if the lesson belongs to the current user
    if lesson.user_id != current_user.id:
        return jsonify({'success': False, 'message': 'Unauthorized access'}), 403

    # Create a temporary file for the presentation
    try:
        with tempfile.NamedTemporaryFile(suffix='.pptx', delete=False) as temp_file:
            temp_filename = temp_file.name

        # Generate the presentation
        create_presentation(lesson.grade_level, lesson.topic, lesson.teaching_strategy,
                          lesson.generated_plan, temp_filename)

        # Create presentation record using Supabase
        presentation = Presentation.create(
            lesson_id=lesson_id,
            file_path=temp_filename
        )

        if not presentation:
            return jsonify({'success': False, 'message': 'Failed to create presentation record'}), 500

        return jsonify({
            'success': True,
            'presentation_id': presentation.id,
            'message': 'Presentation generated successfully'
        })

    except Exception as e:
        return jsonify({'success': False, 'message': f'Error generating presentation: {str(e)}'}), 500

@routes.route('/api/presentations/<int:presentation_id>/download', methods=['GET'])
@login_required
def download_presentation(presentation_id):
    supabase = current_app.config["SUPABASE_CLIENT"]
    
    # Get presentation
    try:
        pres_response = supabase.table('presentations').select("*").eq('id', presentation_id).single().execute()
        if not pres_response.data:
            return jsonify({'success': False, 'message': 'Presentation not found'}), 404
        presentation = Presentation(pres_response.data)
    except:
        return jsonify({'success': False, 'message': 'Presentation not found'}), 404

    # Get lesson
    lesson = Lesson.get_by_id(presentation.lesson_id)
    if not lesson:
        return jsonify({'success': False, 'message': 'Lesson not found'}), 404

    # Check if the presentation belongs to the current user
    if lesson.user_id != current_user.id:
        return jsonify({'success': False, 'message': 'Unauthorized access'}), 403

    # Check if file exists
    if not os.path.exists(presentation.file_path):
        return jsonify({'success': False, 'message': 'Presentation file not found'}), 404

    # Generate filename
    filename = f"Lesson_{lesson.topic.replace(' ', '_')}.pptx"

    # Send the file
    return send_file(
        presentation.file_path,
        as_attachment=True,
        download_name=filename,
        mimetype='application/vnd.openxmlformats-officedocument.presentationml.presentation'
    )

@routes.route('/login', methods=['GET', 'POST'])
def login():
    if current_user.is_authenticated:
        return redirect(url_for('routes.index'))

    language = session.get('language', 'en')
    form = LoginForm()
    
    if form.validate_on_submit():
        user = User.get_by_email(form.email.data)
        if user and user.check_password(form.password.data):
            login_user(user)
            next_page = request.args.get('next')
            if not next_page or urlparse(next_page).netloc != '':
                next_page = url_for('routes.dashboard')
            return redirect(next_page)
        flash('Invalid email or password')

    return render_template('login.html', form=form, language=language)

@routes.route('/register', methods=['GET', 'POST'])
def register():
    if current_user.is_authenticated:
        return redirect(url_for('routes.dashboard'))

    language = session.get('language', 'en')
    form = RegistrationForm()
    
    if form.validate_on_submit():
        user = User.create(
            name=form.name.data,
            email=form.email.data,
            password=form.password.data
        )
        if user:
            flash('Registration successful! You can now log in.')
            return redirect(url_for('routes.login'))
        else:
            flash('Registration failed. Please try again.')

    return render_template('register.html', form=form, language=language)

@routes.route('/dashboard')
@login_required
def dashboard():
    # Get user's lessons using Supabase
    all_lessons = Lesson.get_all_by_user(current_user.id)
    total_lessons = len(all_lessons)
    recent_lessons = all_lessons[:5]  # First 5 (already ordered by date_created desc)

    language = session.get('language', 'en')

    return render_template('dashboard.html',
                          total_lessons=total_lessons,
                          recent_lessons=recent_lessons,
                          language=language)

@routes.route('/lessons')
@login_required
def get_lessons_page():
    # Get all lessons for the current user using Supabase
    lessons = Lesson.get_all_by_user(current_user.id)
    language = session.get('language', 'en')

    return render_template('lesson_list.html', lessons=lessons, language=language)

@routes.route('/logout')
def logout():
    logout_user()
    return redirect(url_for('routes.index'))

@routes.route('/set-language/<language>')
def set_language(language):
    session['language'] = language
    return redirect(request.referrer or url_for('routes.index'))

@routes.route('/create-lesson', methods=['GET', 'POST'])
@login_required
def create_lesson_form():
    language = session.get('language', 'en')
    form = LessonForm()
    
    if form.validate_on_submit():
        if not current_user.deduct_tokens(1, 'lesson_create', 'web'):
            flash('Insufficient tokens' if language == 'en' else 'نفدت التوكنز المتاحة. الرجاء الترقية أو انتظار التجديد الشهري.')
            return render_template('create_lesson.html', form=form, language=language)
        # Generate lesson plan content first
        try:
            lesson_plan = generate_lesson_plan(
                form.grade_level.data,
                form.topic.data,
                form.teaching_strategy.data,
                form.language.data
            )
            gpt_plan = gpt_plans(
                form.grade_level.data,
                form.topic.data,
                form.teaching_strategy.data,
                form.language.data
            )
        except Exception as e:
            flash(f'Error generating lesson plan: {str(e)}')
            return render_template('create_lesson.html', form=form, language=language)

        # Create new lesson using Supabase
        lesson = Lesson.create(
            user_id=current_user.id,
            grade_level=form.grade_level.data,
            topic=form.topic.data,
            teaching_strategy=form.teaching_strategy.data,
            language=form.language.data,
            generated_plan=lesson_plan,
            gpt_plan=gpt_plan
        )

        if lesson:
            flash('Lesson plan created successfully!')
            return redirect(url_for('routes.edit_lesson_form', lesson_id=lesson.id))
        else:
            flash('Failed to create lesson plan')

    return render_template('create_lesson.html', form=form, language=language)

@routes.route('/edit-lesson/<int:lesson_id>', methods=['GET', 'POST'])
@login_required
def edit_lesson_form(lesson_id):
    lesson = Lesson.get_by_id(lesson_id)
    
    if not lesson:
        flash('Lesson not found')
        return redirect(url_for('routes.index'))
    
    if lesson.user_id != current_user.id:
        flash('Unauthorized access')
        return redirect(url_for('routes.index'))

    language = session.get('language', 'en')
    form = EditLessonForm()

    if request.method == 'GET':
        form.generated_plan.data = lesson.generated_plan
        form.gpt_plan.data = lesson.gpt_plan

    if form.validate_on_submit():
        try:
            supabase = current_app.config["SUPABASE_CLIENT"]
            supabase.table('lessons').update({
                'generated_plan': form.generated_plan.data,
                'gpt_plan': form.gpt_plan.data,
                'date_modified': 'now()'
            }).eq('id', lesson_id).execute()
            
            flash('Lesson plan updated successfully!')
            return redirect(url_for('routes.edit_lesson_form', lesson_id=lesson.id))
        except Exception as e:
            flash(f'Error updating lesson: {str(e)}')

    return render_template('edit_lesson.html', form=form, lesson=lesson, language=language)

@routes.route('/user', methods=['GET', 'POST'])
@login_required
def user():
    language = session.get('language', 'en')
    form = UserProfileForm()

    if form.validate_on_submit():
        # Verify current password
        if not current_user.check_password(form.current_password.data):
            flash('Current password is incorrect' if language == 'en' else 'كلمة المرور الحالية غير صحيحة')
            return render_template('user.html', form=form, language=language)

        try:
            supabase = current_app.config["SUPABASE_CLIENT"]
            update_data = {
                'name': form.name.data,
                'email': form.email.data
            }

            # Update password if new password is provided
            if form.new_password.data:
                from werkzeug.security import generate_password_hash
                update_data['password_hash'] = generate_password_hash(form.new_password.data)

            supabase.table('users').update(update_data).eq('id', current_user.id).execute()
            
            flash('Profile updated successfully!' if language == 'en' else 'تم تحديث الملف الشخصي بنجاح!')
            return redirect(url_for('routes.user'))
        except Exception as e:
            flash(f'Error updating profile: {str(e)}')

    # Pre-fill form with current user data
    if request.method == 'GET':
        form.name.data = current_user.name
        form.email.data = current_user.email

    return render_template('user.html', form=form, language=language)

@routes.route('/delete-account', methods=['POST'])
@login_required
def delete_account():
    language = session.get('language', 'en')

    try:
        supabase = current_app.config["SUPABASE_CLIENT"]
        
        # Delete user's lessons first (cascade should handle presentations)
        supabase.table('lessons').delete().eq('user_id', current_user.id).execute()
        
        # Delete user
        supabase.table('users').delete().eq('id', current_user.id).execute()
        
        logout_user()
        flash('Your account has been deleted.' if language == 'en' else 'تم حذف حسابك.')
    except Exception as e:
        flash(f'Error deleting account: {str(e)}')
    
    return redirect(url_for('routes.index'))

@routes.route('/whatsapp-sender', methods=['GET', 'POST'])
@login_required
def whatsapp_sender():
    language = session.get('language', 'en')
    form = WhatsAppMessageForm()

    if form.validate_on_submit():
        try:
            # Process the Excel file
            messages = process_excel_file(
                form.excel_file.data,
                form.student_name_column.data,
                form.grades_column.data,
                form.phone_column.data,
                form.message_template.data
            )

            if not messages:
                flash('No valid data found in the Excel file.' if language == 'en' else 'لم يتم العثور على بيانات صالحة في ملف Excel.', 'warning')
                return render_template('whatsapp_sender.html', form=form, language=language)

            # Store messages in session for the results page
            session['whatsapp_messages'] = messages

            # Open WhatsApp Web for the first message
            if messages:
                open_whatsapp_web(messages[0]['phone'], messages[0]['message'])

            flash(f'Successfully processed {len(messages)} messages. WhatsApp Web has been opened for the first message.' if language == 'en' else
                  f'تمت معالجة {len(messages)} رسالة بنجاح. تم فتح واتساب ويب للرسالة الأولى.', 'success')

            return redirect(url_for('routes.whatsapp_sender'))

        except Exception as e:
            flash(f'Error: {str(e)}' if language == 'en' else f'خطأ: {str(e)}', 'danger')

    return render_template('coningsoon.html', form=form, language=language)

@routes.route('/send-whatsapp/<int:index>', methods=['GET'])
@login_required
def send_individual_whatsapp(index):
    messages = session.get('whatsapp_messages', [])

    if not messages or index >= len(messages):
        flash('Invalid message index or no messages available.', 'danger')
        return redirect(url_for('routes.whatsapp_sender'))

    message = messages[index]
    open_whatsapp_web(message['phone'], message['message'])

    flash(f'WhatsApp Web opened for {message["student_name"]}.', 'success')
    return redirect(url_for('routes.whatsapp_sender'))

@routes.route('/clear-whatsapp-messages', methods=['POST'])
@login_required
def clear_whatsapp_messages():

    if 'whatsapp_messages' in session:
        session.pop('whatsapp_messages')
    return jsonify({'success': True})


@routes.route("/download-lesson/<int:lesson_id>")
@login_required
def download_lesson(lesson_id):

    supabase = current_app.config["SUPABASE_CLIENT"]

    response = (
        supabase.table("lessons")
        .select("*")
        .eq("id", lesson_id)
        .eq("user_id", current_user.id)
        .single()
        .execute()
    )

    lesson = response.data
    if not lesson:
        abort(404)

    doc = Document()

    lines = lesson["generated_plan"].split("\n")

    for line in lines:
        line = line.strip()

        # عنوان رئيسي ###
        if line.startswith("###"):
            doc.add_heading(line.replace("###", "").strip(), level=1)

        # عنوان فرعي ####
        elif line.startswith("####"):
            doc.add_heading(line.replace("####", "").strip(), level=2)

        # فاصل ---
        elif line.startswith("---"):
            doc.add_page_break()

        # Bullet points
        elif line.startswith("- "):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")

        # نص عادي
        elif line:
            p = doc.add_paragraph()
            run = p.add_run(line.replace("**", ""))
            if "**" in line:
                run.bold = True

    file_stream = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(file_stream.name)

    return send_file(
        file_stream.name,
        as_attachment=True,
        download_name=f'{lesson["topic"]}.docx',
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
